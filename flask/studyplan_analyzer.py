from __future__ import annotations

import base64
import csv
import json
import mimetypes
import os
import re
import uuid
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Iterable

import requests
from openpyxl import Workbook
from openpyxl.styles import Font

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

from excel_layout_builder import build_structured_study_plan_workbook


# =========================
# Configuration
# =========================

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
GEMINI_TIMEOUT = int(os.getenv("GEMINI_TIMEOUT", "120"))

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".json",
    ".csv",
    ".xlsx",
    ".xls",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}


# =========================
# Data model
# =========================

@dataclass
class CourseRecord:
    course_code: str = ""
    course_name: str = ""
    credit_hours: int | None = None
    prerequisites: list[str] | list[dict[str, Any]] | None = None
    year_no: int | None = None
    semester_no: int | None = None
    category: str = ""
    notes: str = ""

    def __post_init__(self) -> None:
        if self.prerequisites is None:
            self.prerequisites = []


# =========================
# Normalization helpers
# =========================

COURSE_CODE_RE = re.compile(r"\b([A-Z]{2,6}\s*[-_]?\s*\d{2,4}[A-Z]?)\b", re.IGNORECASE)
PLACEHOLDER_CODE_RE = re.compile(r"\b([A-Z]{2,8}X{2,6})\b", re.IGNORECASE)


def normalize_space(text: Any) -> str:
    if text is None:
        return ""
    return re.sub(r"\s+", " ", str(text)).strip()


def normalize_course_code(value: Any) -> str:
    text = normalize_space(value).upper().replace("_", "").replace("-", "").replace(" ", "")
    if not text:
        return ""

    placeholder_match = PLACEHOLDER_CODE_RE.search(text)
    if placeholder_match:
        return placeholder_match.group(1).upper()

    code_match = re.search(r"([A-Z]{2,6}\d{2,4}[A-Z]?)", text)
    if code_match:
        return code_match.group(1).upper()

    return ""


def normalize_course_name(value: Any) -> str:
    text = normalize_space(value)
    if not text:
        return ""

    text = re.sub(r"\b(pre[- ]?reqs?|prerequisites?)\b.*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(co[- ]?reqs?|corequisites?)\b.*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bcredits?\b.*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^\s*[A-Z]{2,6}\s*[-_]?\s*\d{2,4}[A-Z]?\s*", "", text, flags=re.IGNORECASE)

    return normalize_space(text)


def normalize_credit_hours(value: Any) -> int | None:
    if value is None:
        return None

    if isinstance(value, (int, float)):
        ivalue = int(value)
        return ivalue if 0 <= ivalue <= 10 else None

    text = normalize_space(value)
    if not text:
        return None

    match = re.search(r"\b(\d+(?:\.\d+)?)\b", text)
    if not match:
        return None

    try:
        ivalue = int(float(match.group(1)))
    except ValueError:
        return None

    return ivalue if 0 <= ivalue <= 10 else None


def normalize_prerequisites(value: Any) -> list[str]:
    if value is None:
        return []

    if isinstance(value, list):
        items = value
    else:
        text = normalize_space(value)
        if not text:
            return []
        items = re.split(r"[;,/]|(?:\band\b)|(?:\bor\b)", text, flags=re.IGNORECASE)

    result: list[str] = []
    for item in items:
        code = normalize_course_code(item)
        if code and code not in result:
            result.append(code)
    return result


def normalize_int(value: Any) -> int | None:
    if value is None:
        return None
    if isinstance(value, int):
        return value
    text = normalize_space(value)
    if not text:
        return None
    match = re.search(r"\d+", text)
    return int(match.group()) if match else None


def is_reasonable_course_row(course: CourseRecord) -> bool:
    if not course.course_code and not course.course_name:
        return False
    if course.credit_hours is not None and not (0 <= course.credit_hours <= 10):
        return False
    if len(course.course_name) > 180:
        return False
    return True


def normalize_course_record(raw: dict[str, Any]) -> CourseRecord:
    return CourseRecord(
        course_code=normalize_course_code(raw.get("course_code") or raw.get("code")),
        course_name=normalize_course_name(raw.get("course_name") or raw.get("name") or raw.get("title")),
        credit_hours=normalize_credit_hours(raw.get("credit_hours") or raw.get("credits") or raw.get("credit")),
        prerequisites=normalize_prerequisites(raw.get("prerequisites") or raw.get("prereqs") or raw.get("prerequisite")),
        year_no=normalize_int(raw.get("year_no") or raw.get("year")),
        semester_no=normalize_int(raw.get("semester_no") or raw.get("semester") or raw.get("term")),
        category=normalize_space(raw.get("category")),
        notes=normalize_space(raw.get("notes")),
    )


# =========================
# Generic helpers
# =========================

def _safe_str(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _safe_int(value: Any) -> int | None:
    if value is None or value == "":
        return None
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return None


# =========================
# File type detection
# =========================

def detect_file_type(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")
    return suffix.lstrip(".")


def guess_mime_type(file_path: str) -> str:
    mime_type, _ = mimetypes.guess_type(file_path)
    return mime_type or "application/octet-stream"


# =========================
# Structured file readers
# =========================

def read_text_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_json_file(file_path: str) -> Any:
    with open(file_path, "r", encoding="utf-8") as f:
        return json.load(f)


def read_csv_rows(file_path: str) -> list[dict[str, Any]]:
    with open(file_path, "r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.DictReader(f)
        return [dict(row) for row in reader]


def read_xlsx_rows(file_path: str) -> list[dict[str, Any]]:
    if load_workbook is None:
        raise RuntimeError("openpyxl is not installed.")

    wb = load_workbook(file_path, data_only=True)
    rows_out: list[dict[str, Any]] = []

    for ws in wb.worksheets:
        raw_rows = list(ws.iter_rows(values_only=True))
        if not raw_rows:
            continue

        headers = [normalize_space(x) for x in raw_rows[0]]
        if not any(headers):
            continue

        for row in raw_rows[1:]:
            values = list(row)
            row_dict = {}
            for idx, header in enumerate(headers):
                if not header:
                    continue
                row_dict[header] = values[idx] if idx < len(values) else None
            if any(v not in (None, "") for v in row_dict.values()):
                rows_out.append(row_dict)

    return rows_out


def read_docx_text(file_path: str) -> str:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")

    doc = Document(file_path)
    parts: list[str] = []

    for p in doc.paragraphs:
        text = normalize_space(p.text)
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_space(cell.text) for cell in row.cells if normalize_space(cell.text)]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_pdf_text(file_path: str) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed.")

    reader = PdfReader(file_path)
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = normalize_space(text)
        if text:
            parts.append(text)
    return "\n".join(parts)


# =========================
# Local fallback parsing
# =========================

def parse_rows_from_generic_records(records: Iterable[dict[str, Any]]) -> list[CourseRecord]:
    normalized: list[CourseRecord] = []
    for raw in records:
        lowered = {str(k).strip().lower(): v for k, v in raw.items()}

        mapped = {
            "course_code": lowered.get("course_code") or lowered.get("code") or lowered.get("course") or lowered.get("course id"),
            "course_name": lowered.get("course_name") or lowered.get("name") or lowered.get("title") or lowered.get("course title"),
            "credit_hours": lowered.get("credit_hours") or lowered.get("credits") or lowered.get("credit") or lowered.get("cr"),
            "prerequisites": lowered.get("prerequisites") or lowered.get("prereqs") or lowered.get("prerequisite"),
            "year_no": lowered.get("year_no") or lowered.get("year"),
            "semester_no": lowered.get("semester_no") or lowered.get("semester") or lowered.get("term"),
            "category": lowered.get("category") or lowered.get("type"),
            "notes": lowered.get("notes"),
        }

        course = normalize_course_record(mapped)
        if is_reasonable_course_row(course):
            normalized.append(course)

    return normalized


def parse_courses_from_text(text: str) -> list[CourseRecord]:
    lines = [normalize_space(line) for line in text.splitlines()]
    lines = [line for line in lines if line]

    courses: list[CourseRecord] = []

    for line in lines:
        code_match = COURSE_CODE_RE.search(line)
        placeholder_match = PLACEHOLDER_CODE_RE.search(line)

        code = ""
        if code_match:
            code = normalize_course_code(code_match.group(1))
        elif placeholder_match:
            code = normalize_course_code(placeholder_match.group(1))

        if not code:
            continue

        credit = normalize_credit_hours(line)
        prereqs = normalize_prerequisites(line)

        name = line
        if code:
            name = re.sub(COURSE_CODE_RE, "", name, count=1)
            name = re.sub(PLACEHOLDER_CODE_RE, "", name, count=1)

        name = re.sub(r"\b(?:year|semester|term)\s*\d+\b", "", name, flags=re.IGNORECASE)
        name = re.sub(r"\b\d+(?:\.\d+)?\b", "", name)
        name = normalize_course_name(name)

        course = CourseRecord(
            course_code=code,
            course_name=name,
            credit_hours=credit,
            prerequisites=prereqs,
        )

        if is_reasonable_course_row(course):
            courses.append(course)

    deduped: dict[str, CourseRecord] = {}
    for course in courses:
        key = course.course_code or course.course_name
        if not key:
            continue

        current = deduped.get(key)
        if current is None:
            deduped[key] = course
            continue

        current_score = int(bool(current.course_name)) + int(current.credit_hours is not None) + len(current.prerequisites)
        new_score = int(bool(course.course_name)) + int(course.credit_hours is not None) + len(course.prerequisites)
        if new_score > current_score:
            deduped[key] = course

    return list(deduped.values())


# =========================
# Gemini extraction
# =========================

def gemini_prompt(program_hint: str | None = None) -> str:
    hint = f"Program hint: {program_hint}\n" if program_hint else ""
    return f"""
You are a university study plan extraction engine.

Read the uploaded study plan document and extract only the study plan structure.
Return valid JSON only.
Do not add markdown.
Do not add explanations.

{hint}
Rules:
1. Detect course boundaries carefully.
2. Extract one row per course.
3. Separate:
   - course_code
   - course_name
   - credit_hours
   - prerequisites
   - year_no
   - semester_no
   - category
4. Do not merge prerequisites or neighboring rows into course_name.
5. If prerequisites are free-text rules, keep only real course codes in prerequisites.
6. Preserve elective placeholders like GHALXXX, GSOSXXX, C3SXXX, AIXXX.
7. If a field is unknown, return null or empty string/list.
8. credit_hours must be numeric and realistic per course.

Return exactly this JSON shape:
{{
  "program_name": "",
  "catalog_year": "",
  "courses": [
    {{
      "course_code": "",
      "course_name": "",
      "credit_hours": null,
      "prerequisites": [],
      "year_no": null,
      "semester_no": null,
      "category": "",
      "notes": ""
    }}
  ],
  "slot_rules": {{}}
}}
""".strip()

import time

def call_gemini_with_file(file_path: str, prompt: str, max_retries: int = 4, backoff_seconds: float = 2.0) -> dict[str, Any]:
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not set.")

    with open(file_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"

    payload = {
        "contents": [
            {
                "parts": [
                    {"text": prompt},
                    {
                        "inline_data": {
                            "mime_type": guess_mime_type(file_path),
                            "data": encoded,
                        }
                    },
                ]
            }
        ],
        "generationConfig": {
            "temperature": 0,
            "response_mime_type": "application/json",
        },
    }

    last_error = None

    for attempt in range(1, max_retries + 1):
        try:
            response = requests.post(url, json=payload, timeout=GEMINI_TIMEOUT)

            if response.status_code in {429, 500, 502, 503, 504}:
                raise RuntimeError(f"Gemini temporary error {response.status_code}: {response.text[:500]}")

            response.raise_for_status()
            data = response.json()

            try:
                text = data["candidates"][0]["content"]["parts"][0]["text"]
            except (KeyError, IndexError) as exc:
                raise RuntimeError(f"Unexpected Gemini response: {data}") from exc

            try:
                return json.loads(text)
            except json.JSONDecodeError as exc:
                raise RuntimeError(f"Gemini did not return valid JSON: {text[:1000]}") from exc

        except Exception as exc:
            last_error = exc
            if attempt == max_retries:
                break
            time.sleep(backoff_seconds * attempt)

    raise RuntimeError(f"Gemini file call failed after {max_retries} attempts: {last_error}")


import time

def _call_gemini_json(prompt: str, max_retries: int = 4, backoff_seconds: float = 2.0) -> Any:
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not set.")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"

    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0,
            "response_mime_type": "application/json",
        },
    }

    last_error = None

    for attempt in range(1, max_retries + 1):
        try:
            response = requests.post(url, json=payload, timeout=GEMINI_TIMEOUT)

            if response.status_code in {429, 500, 502, 503, 504}:
                raise RuntimeError(f"Gemini temporary error {response.status_code}: {response.text[:500]}")

            response.raise_for_status()
            data = response.json()

            try:
                text = data["candidates"][0]["content"]["parts"][0]["text"]
            except (KeyError, IndexError) as exc:
                raise RuntimeError(f"Unexpected Gemini response: {data}") from exc

            try:
                return json.loads(text)
            except json.JSONDecodeError as exc:
                raise RuntimeError(f"Gemini returned invalid JSON: {text[:1000]}") from exc

        except Exception as exc:
            last_error = exc
            if attempt == max_retries:
                break
            time.sleep(backoff_seconds * attempt)

    raise RuntimeError(f"Gemini JSON call failed after {max_retries} attempts: {last_error}")


    
def extract_study_plan_with_gemini(file_path: str, program_hint: str | None = None) -> dict[str, Any] | None:
    payload = call_gemini_with_file(file_path, gemini_prompt(program_hint))

    raw_courses = payload.get("courses", [])
    courses: list[CourseRecord] = []

    for raw in raw_courses:
        course = normalize_course_record(raw)
        if is_reasonable_course_row(course):
            courses.append(course)

    if not courses:
        return None

    return {
        "program_name": normalize_space(payload.get("program_name")) or normalize_space(program_hint),
        "catalog_year": normalize_space(payload.get("catalog_year")),
        "courses": [asdict(course) for course in courses],
        "slot_rules": payload.get("slot_rules", {}) or {},
        "source_type": "gemini",
    }


# =========================
# Study plan main public function
# =========================
def analyze_study_plan(file_path: str, program_hint: str | None = None) -> dict[str, Any]:
    file_type = detect_file_type(file_path)

    if file_type == "json":
        data = read_json_file(file_path)

        if isinstance(data, dict) and "courses" in data:
            raw_courses = data.get("courses", [])
            courses = [normalize_course_record(row) for row in raw_courses]
            courses = [c for c in courses if is_reasonable_course_row(c)]
            return {
                "program_name": normalize_space(data.get("program_name")) or normalize_space(program_hint),
                "catalog_year": normalize_space(data.get("catalog_year")),
                "courses": [asdict(c) for c in courses],
                "slot_rules": data.get("slot_rules", {}) or {},
                "source_type": "json",
            }

        if isinstance(data, list):
            courses = parse_rows_from_generic_records(data)
            return {
                "program_name": normalize_space(program_hint),
                "catalog_year": "",
                "courses": [asdict(c) for c in courses],
                "slot_rules": {},
                "source_type": "json",
            }

        raise ValueError("Unsupported JSON structure.")

    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not set. Gemini-only study-plan extraction cannot run.")

    gemini_result = extract_study_plan_with_gemini(file_path, program_hint=program_hint)
    if not gemini_result or not gemini_result.get("courses"):
        raise RuntimeError("Gemini study-plan extraction returned no courses.")

    return gemini_result
# =========================
# Comparison helpers
# =========================

def _prepare_plan_courses_for_model(plan_data: dict[str, Any]) -> list[dict[str, Any]]:
    courses = plan_data.get("courses", []) or []
    prepared: list[dict[str, Any]] = []

    for idx, c in enumerate(courses, start=1):
        prepared.append(
            {
                "plan_index": idx,
                "course_code": normalize_course_code(c.get("course_code")),
                "course_name": normalize_course_name(c.get("course_name")),
                "credit_hours": _safe_int(c.get("credit_hours")),
                "prerequisites": normalize_prerequisites(c.get("prerequisites")),
                "year_no": _safe_int(c.get("year_no")),
                "semester_no": _safe_int(c.get("semester_no")),
                "category": _safe_str(c.get("category")),
                "notes": _safe_str(c.get("notes")),
            }
        )

    return prepared


def _prepare_transcript_courses_for_model(transcript_data: dict[str, Any]) -> list[dict[str, Any]]:
    courses = transcript_data.get("courses", []) or []
    prepared: list[dict[str, Any]] = []

    for idx, c in enumerate(courses, start=1):
        prepared.append(
            {
                "transcript_index": idx,
                "course_code": normalize_course_code(c.get("course_code")),
                "course_name": normalize_course_name(c.get("course_name")),
                "credit_hours": _safe_int(c.get("credit_hours")),
                "grade": _safe_str(c.get("grade")),
                "status": _safe_str(c.get("status")).lower(),
                "term_taken": _safe_str(c.get("term_taken")),
                "notes": _safe_str(c.get("notes")),
                "points": c.get("points", 0),
            }
        )

    return prepared


def _completed_transcript_codes(transcript_courses: list[dict[str, Any]]) -> set[str]:
    completed = set()
    for c in transcript_courses:
        if _safe_str(c.get("status")).lower() == "completed":
            code = normalize_course_code(c.get("course_code"))
            if code:
                completed.add(code)
    return completed


def _model_match_prompt(
    plan_courses: list[dict[str, Any]],
    transcript_courses: list[dict[str, Any]],
) -> str:
    return f"""
You are a university academic-audit matching engine.

Your task is to compare:
1. study plan courses
2. transcript courses

and return a JSON array where EACH study plan course gets exactly one merged result row.

Strict matching rules:
1. Prefer exact course code match first.
2. If course code is absent in the study plan row, you may use course-name similarity.
3. Do NOT match two different study plan rows to one transcript row unless the row is an elective placeholder or generic slot.
4. Elective placeholders like AIXXX, GHALXXX, GSOSXXX, C3SXXX are slots, not exact real transcript course codes.
5. A transcript course with status "completed" means completed.
6. A transcript course with status "in_progress" means in progress.
7. If no transcript course matches, set status to "not_completed".
8. If prerequisites are listed and the course is not completed/in progress, determine whether it is blocked:
   - blocked = true only if at least one prerequisite course code is missing from completed transcript course codes.
9. Do not invent courses.
10. Return valid JSON only. No markdown. No explanation.

Return exactly this array schema:
[
  {{
    "plan_index": 1,
    "study_plan_course_code": "",
    "study_plan_course_name": "",
    "study_plan_credit_hours": null,
    "year_no": null,
    "semester_no": null,
    "category": "",
    "prerequisites": [],
    "matched": false,
    "matched_transcript_index": null,
    "matched_transcript_course_code": "",
    "matched_transcript_course_name": "",
    "matched_transcript_credit_hours": null,
    "grade": "",
    "term_taken": "",
    "status": "not_completed",
    "match_type": "none",
    "blocked_by_prerequisite": false,
    "notes": ""
  }}
]

Allowed match_type values:
- exact_code
- normalized_code
- name_similarity
- elective_slot
- none

Study plan courses JSON:
{json.dumps(plan_courses, ensure_ascii=False, indent=2)}

Transcript courses JSON:
{json.dumps(transcript_courses, ensure_ascii=False, indent=2)}
""".strip()


def _postprocess_model_merged_rows(
    model_rows: list[dict[str, Any]],
    plan_courses: list[dict[str, Any]],
    transcript_courses: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    plan_index_map = {row["plan_index"]: row for row in plan_courses}
    transcript_index_map = {row["transcript_index"]: row for row in transcript_courses}
    completed_codes = _completed_transcript_codes(transcript_courses)

    processed: list[dict[str, Any]] = []

    for row in model_rows:
        plan_index = _safe_int(row.get("plan_index"))
        if not plan_index or plan_index not in plan_index_map:
            continue

        plan_row = plan_index_map[plan_index]
        transcript_index = _safe_int(row.get("matched_transcript_index"))
        transcript_row = transcript_index_map.get(transcript_index) if transcript_index else None

        prerequisites = row.get("prerequisites")
        if not isinstance(prerequisites, list):
            prerequisites = plan_row.get("prerequisites", []) or []

        blocked = bool(row.get("blocked_by_prerequisite"))
        status = _safe_str(row.get("status")).lower() or "not_completed"

        if status not in {"completed", "in_progress", "not_completed"}:
            status = "not_completed"

        if status == "not_completed":
            prereq_codes = [normalize_course_code(p) for p in prerequisites if normalize_course_code(p)]
            if prereq_codes:
                blocked = any(code not in completed_codes for code in prereq_codes)

        merged = {
            "plan_index": plan_index,
            "course_code": _safe_str(row.get("study_plan_course_code")) or plan_row.get("course_code", ""),
            "course_name": _safe_str(row.get("study_plan_course_name")) or plan_row.get("course_name", ""),
            "credit_hours": _safe_int(row.get("study_plan_credit_hours"))
            if row.get("study_plan_credit_hours") is not None
            else plan_row.get("credit_hours"),
            "year_no": _safe_int(row.get("year_no"))
            if row.get("year_no") is not None
            else plan_row.get("year_no"),
            "semester_no": _safe_int(row.get("semester_no"))
            if row.get("semester_no") is not None
            else plan_row.get("semester_no"),
            "category": _safe_str(row.get("category")) or plan_row.get("category", ""),
            "prerequisites": prerequisites,
            "status": status,
            "blocked_by_prerequisite": blocked,
            "matched": bool(row.get("matched")),
            "match_type": _safe_str(row.get("match_type")) or "none",
            "transcript_course_code": transcript_row.get("course_code", "") if transcript_row else _safe_str(row.get("matched_transcript_course_code")),
            "transcript_course_name": transcript_row.get("course_name", "") if transcript_row else _safe_str(row.get("matched_transcript_course_name")),
            "transcript_credit_hours": transcript_row.get("credit_hours") if transcript_row else _safe_int(row.get("matched_transcript_credit_hours")),
            "grade": transcript_row.get("grade", "") if transcript_row else _safe_str(row.get("grade")),
            "term_taken": transcript_row.get("term_taken", "") if transcript_row else _safe_str(row.get("term_taken")),
            "notes": _safe_str(row.get("notes")) or plan_row.get("notes", ""),
        }

        processed.append(merged)

    existing_plan_indices = {row["plan_index"] for row in processed}
    for plan_row in plan_courses:
        if plan_row["plan_index"] in existing_plan_indices:
            continue

        prereqs = plan_row.get("prerequisites", []) or []
        prereq_codes = [normalize_course_code(p) for p in prereqs if normalize_course_code(p)]
        blocked = any(code not in completed_codes for code in prereq_codes) if prereq_codes else False

        processed.append(
            {
                "plan_index": plan_row["plan_index"],
                "course_code": plan_row.get("course_code", ""),
                "course_name": plan_row.get("course_name", ""),
                "credit_hours": plan_row.get("credit_hours"),
                "year_no": plan_row.get("year_no"),
                "semester_no": plan_row.get("semester_no"),
                "category": plan_row.get("category", ""),
                "prerequisites": prereqs,
                "status": "not_completed",
                "blocked_by_prerequisite": blocked,
                "matched": False,
                "match_type": "none",
                "transcript_course_code": "",
                "transcript_course_name": "",
                "transcript_credit_hours": None,
                "grade": "",
                "term_taken": "",
                "notes": plan_row.get("notes", ""),
            }
        )

    processed.sort(key=lambda x: (x.get("year_no") or 999, x.get("semester_no") or 999, x.get("plan_index") or 999))
    return processed


def _match_study_plan_courses_with_model(
    plan_data: dict[str, Any],
    transcript_data: dict[str, Any],
) -> list[dict[str, Any]]:
    plan_courses = _prepare_plan_courses_for_model(plan_data)
    transcript_courses = _prepare_transcript_courses_for_model(transcript_data)

    if not plan_courses:
        return []

    if not transcript_courses:
        return _postprocess_model_merged_rows([], plan_courses, transcript_courses)

    prompt = _model_match_prompt(plan_courses, transcript_courses)
    model_output = _call_gemini_json(prompt)

    if not isinstance(model_output, list):
        raise RuntimeError("Gemini comparison output must be a list.")

    return _postprocess_model_merged_rows(model_output, plan_courses, transcript_courses)


# =========================
# Rule-based fallback matcher
# =========================

def _build_transcript_lookup(transcript_data: dict[str, Any]) -> tuple[dict[str, dict[str, Any]], dict[str, list[dict[str, Any]]]]:
    by_code: dict[str, dict[str, Any]] = {}
    by_name: dict[str, list[dict[str, Any]]] = {}

    for c in transcript_data.get("courses", []) or []:
        code = normalize_course_code(c.get("course_code"))
        name = normalize_course_name(c.get("course_name")).lower()

        if code:
            current = by_code.get(code)
            if current is None:
                by_code[code] = c
            else:
                current_status = _safe_str(current.get("status")).lower()
                new_status = _safe_str(c.get("status")).lower()
                priority = {"completed": 3, "in_progress": 2, "not_taken": 1, "": 0}
                if priority.get(new_status, 0) > priority.get(current_status, 0):
                    by_code[code] = c

        if name:
            by_name.setdefault(name, []).append(c)

    return by_code, by_name


def _prereq_codes_satisfied(prerequisites: list[str], satisfied_codes: set[str]) -> bool:
    prereq_codes = [normalize_course_code(p) for p in prerequisites if normalize_course_code(p)]
    return all(code in satisfied_codes for code in prereq_codes)


def _canonical_merged_row(plan_course: dict[str, Any], transcript_course: dict[str, Any] | None, satisfied_codes: set[str]) -> dict[str, Any]:
    status = "not_completed"
    grade = ""
    term_taken = ""
    transcript_course_code = ""
    transcript_course_name = ""
    transcript_credit_hours = None
    matched = False
    match_type = "none"

    if transcript_course:
        matched = True
        transcript_course_code = _safe_str(transcript_course.get("course_code"))
        transcript_course_name = _safe_str(transcript_course.get("course_name"))
        transcript_credit_hours = _safe_int(transcript_course.get("credit_hours"))
        grade = _safe_str(transcript_course.get("grade"))
        term_taken = _safe_str(transcript_course.get("term_taken"))
        status = _safe_str(transcript_course.get("status")).lower() or "not_completed"
        if status not in {"completed", "in_progress", "not_completed"}:
            status = "not_completed"
        match_type = "exact_code" if normalize_course_code(plan_course.get("course_code")) else "name_similarity"

    prerequisites = normalize_prerequisites(plan_course.get("prerequisites"))
    blocked = False
    if status == "not_completed" and prerequisites:
        blocked = not _prereq_codes_satisfied(prerequisites, satisfied_codes)

    return {
        "course_code": _safe_str(plan_course.get("course_code")),
        "course_name": _safe_str(plan_course.get("course_name")),
        "credit_hours": _safe_int(plan_course.get("credit_hours")),
        "year_no": _safe_int(plan_course.get("year_no")),
        "semester_no": _safe_int(plan_course.get("semester_no")),
        "category": _safe_str(plan_course.get("category")),
        "prerequisites": prerequisites,
        "status": status,
        "blocked_by_prerequisite": blocked,
        "matched": matched,
        "match_type": match_type,
        "transcript_course_code": transcript_course_code,
        "transcript_course_name": transcript_course_name,
        "transcript_credit_hours": transcript_credit_hours,
        "grade": grade,
        "term_taken": term_taken,
        "notes": _safe_str(plan_course.get("notes")),
    }


def _match_study_plan_courses(plan_data: dict[str, Any], transcript_data: dict[str, Any]) -> list[dict[str, Any]]:
    by_code, by_name = _build_transcript_lookup(transcript_data)

    satisfied_codes = {
        normalize_course_code(c.get("course_code"))
        for c in transcript_data.get("courses", []) or []
        if _safe_str(c.get("status")).lower() == "completed" and normalize_course_code(c.get("course_code"))
    }

    merged_rows: list[dict[str, Any]] = []
    for plan_course in plan_data.get("courses", []) or []:
        plan_code = normalize_course_code(plan_course.get("course_code"))
        plan_name = normalize_course_name(plan_course.get("course_name")).lower()

        transcript_course = None
        if plan_code:
            transcript_course = by_code.get(plan_code)
        elif plan_name and by_name.get(plan_name):
            transcript_course = by_name[plan_name][0]

        merged_rows.append(_canonical_merged_row(plan_course, transcript_course, satisfied_codes))

    return merged_rows


# =========================
# Combined analysis pipeline
# =========================

def _compute_summary(merged_rows: list[dict[str, Any]]) -> dict[str, Any]:
    total = len(merged_rows)
    completed = sum(1 for r in merged_rows if r.get("status") == "completed")
    in_progress = sum(1 for r in merged_rows if r.get("status") == "in_progress")
    blocked = sum(1 for r in merged_rows if r.get("status") == "not_completed" and r.get("blocked_by_prerequisite"))
    remaining = sum(1 for r in merged_rows if r.get("status") == "not_completed")

    return {
        "total_courses": total,
        "completed_courses": completed,
        "in_progress_courses": in_progress,
        "remaining_courses": remaining,
        "blocked_courses": blocked,
    }


def _eligible_next_semester(merged_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    eligible = [
        r for r in merged_rows
        if r.get("status") == "not_completed" and not r.get("blocked_by_prerequisite")
    ]
    eligible.sort(key=lambda x: (x.get("year_no") or 999, x.get("semester_no") or 999, x.get("course_code") or ""))
    return eligible[:8]


def _generate_advice(merged_rows: list[dict[str, Any]], summary: dict[str, Any]) -> list[str]:
    advice: list[str] = []

    if summary.get("blocked_courses", 0) > 0:
        advice.append("Some remaining courses are currently blocked by unmet prerequisites.")

    if summary.get("in_progress_courses", 0) > 0:
        advice.append("You have courses in progress. Re-run the audit after grades are finalized.")

    eligible = _eligible_next_semester(merged_rows)
    if eligible:
        codes = [r.get("course_code") or r.get("course_name") for r in eligible[:5]]
        advice.append(f"Priority candidates for upcoming registration: {', '.join(codes)}.")

    if summary.get("remaining_courses", 0) == 0:
        advice.append("All study-plan courses appear completed or currently in progress.")

    return advice


def _build_audit_workbook(merged_rows: list[dict[str, Any]], output_path: str) -> str:
    wb = Workbook()
    ws = wb.active
    ws.title = "Study Plan Audit"

    headers = [
        "Year",
        "Semester",
        "Course Code",
        "Course Name",
        "Credit Hours",
        "Prerequisites",
        "Status",
        "Grade",
        "Term Taken",
        "Matched Transcript Code",
        "Notes",
    ]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)

    for row in merged_rows:
        ws.append(
            [
                row.get("year_no"),
                row.get("semester_no"),
                row.get("course_code"),
                row.get("course_name"),
                row.get("credit_hours"),
                ", ".join(row.get("prerequisites") or []),
                row.get("status"),
                row.get("grade"),
                row.get("term_taken"),
                row.get("transcript_course_code"),
                row.get("notes"),
            ]
        )

    for column in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in column)
        ws.column_dimensions[column[0].column_letter].width = min(max(max_length + 2, 12), 48)

    os.makedirs(os.path.dirname(output_path), exist_ok=True) if os.path.dirname(output_path) else None
    wb.save(output_path)
    return output_path


def _infer_total_required_credits(merged_rows: list[dict[str, Any]], fallback: int = 132) -> int:
    total = 0
    for row in merged_rows:
        try:
            total += int(float(row.get("credit_hours") or 0))
        except (TypeError, ValueError):
            continue
    return total or fallback

def analyze_transcript_and_study_plan_data(
    transcript_data: dict[str, Any],
    plan_data: dict[str, Any],
    use_model_comparison: bool = True,
    output_dir: str | None = None,
) -> dict[str, Any]:
    if not use_model_comparison:
        raise RuntimeError("Model comparison is required, but use_model_comparison=False was passed.")

    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not set. Gemini-only comparison cannot run.")

    merged_rows = _match_study_plan_courses_with_model(plan_data, transcript_data)
    summary = _compute_summary(merged_rows)
    advice = _generate_advice(merged_rows, summary)
    audit_excel_path = ""
    audit_excel_filename = ""
    structured_excel_path = ""
    structured_excel_filename = ""

    if output_dir:
        audit_excel_path = _build_audit_workbook(
            merged_rows,
            os.path.join(output_dir, f"study_plan_audit_{uuid.uuid4().hex}.xlsx"),
        )
        audit_excel_filename = os.path.basename(audit_excel_path)

        structured_excel_path = build_structured_study_plan_workbook(
            merged_rows=merged_rows,
            output_path=os.path.join(output_dir, f"structured_study_plan_{uuid.uuid4().hex}.xlsx"),
            program_name=plan_data.get("program_name") or "Study Plan",
            total_required_credits=_infer_total_required_credits(merged_rows),
        )
        structured_excel_filename = os.path.basename(structured_excel_path)

    return {
        "student": transcript_data.get("student", {}) or {},
        "study_plan_meta": {
            "program_name": plan_data.get("program_name", ""),
            "catalog_year": plan_data.get("catalog_year", ""),
            "source_type": plan_data.get("source_type", ""),
        },
        "comparison_engine": "gemini",
        "summary": summary,
        "advice": advice,
        "merged_rows": merged_rows,
        "preview_rows": merged_rows[:10],
        "excel_path": audit_excel_path,
        "excel_filename": audit_excel_filename,
        "structured_excel_path": structured_excel_path,
        "structured_excel_filename": structured_excel_filename,
    }
def analyze_transcript_and_study_plan(*args: Any, **kwargs: Any) -> dict[str, Any]:
    """
    Full pipeline:
    - extract transcript
    - extract study plan
    - compare them
    - generate Excel outputs

    Supports both:
    - transcript_file_path / study_plan_file_path
    - transcript_path / study_plan_path
    """

    from pdf_extractor import extract_transcript_data
    import os
    import uuid

    transcript_file_path = kwargs.pop("transcript_file_path", None)
    study_plan_file_path = kwargs.pop("study_plan_file_path", None)
    program_hint = kwargs.pop("program_hint", None)
    use_model_comparison = kwargs.pop("use_model_comparison", True)
    output_dir = kwargs.pop("output_dir", None)
    transcript_path = kwargs.pop("transcript_path", None)
    study_plan_path = kwargs.pop("study_plan_path", None)

    if len(args) >= 1 and transcript_file_path is None:
        transcript_file_path = args[0]
    if len(args) >= 2 and study_plan_file_path is None:
        study_plan_file_path = args[1]
    if len(args) >= 3 and program_hint is None:
        program_hint = args[2]
    if len(args) >= 4:
        use_model_comparison = args[3]
    if len(args) >= 5 and output_dir is None:
        output_dir = args[4]
    if len(args) > 5:
        raise TypeError("analyze_transcript_and_study_plan() accepts at most 5 positional arguments.")
    if kwargs:
        unexpected = ", ".join(sorted(kwargs.keys()))
        raise TypeError(f"Unexpected keyword argument(s): {unexpected}")

    transcript_file_path = transcript_file_path or transcript_path
    study_plan_file_path = study_plan_file_path or study_plan_path

    if not transcript_file_path:
        raise ValueError("transcript_file_path is required.")
    if not study_plan_file_path:
        raise ValueError("study_plan_file_path is required.")

    print("[1/4] Extracting transcript data...")
    transcript_data = extract_transcript_data(transcript_file_path)
    print(f"[1/4] Done. Transcript courses: {len(transcript_data.get('courses', []))}")

    print("[2/4] Extracting study plan with Gemini...")
    plan_data = analyze_study_plan(study_plan_file_path, program_hint=program_hint)
    print(f"[2/4] Done. Study plan courses: {len(plan_data.get('courses', []))}")

    print("[3/4] Comparing study plan and transcript with Gemini...")
    result = analyze_transcript_and_study_plan_data(
        transcript_data=transcript_data,
        plan_data=plan_data,
        use_model_comparison=use_model_comparison,
        output_dir=output_dir,
    )
    print("[3/4] Done.")

    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(__file__), "uploads")
    os.makedirs(output_dir, exist_ok=True)

    # audit excel
    audit_filename = f"study_plan_audit_{uuid.uuid4().hex}.xlsx"
    audit_path = os.path.join(output_dir, audit_filename)
    _build_audit_workbook(result["merged_rows"], audit_path)

    # structured excel
    structured_filename = f"structured_study_plan_{uuid.uuid4().hex}.xlsx"
    structured_path = os.path.join(output_dir, structured_filename)

    build_structured_study_plan_workbook(
        merged_rows=result["merged_rows"],
        output_path=structured_path,
        program_name=result.get("study_plan_meta", {}).get("program_name", "Study Plan"),
        total_required_credits=_infer_total_required_credits(result["merged_rows"], fallback=132),
    )

    result["excel_path"] = audit_path
    result["excel_filename"] = audit_filename
    result["structured_excel_path"] = structured_path
    result["structured_excel_filename"] = structured_filename

    return result
# =========================
# CLI usage
# =========================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Analyze a study plan or compare a study plan with a transcript.")
    parser.add_argument("study_plan_file", help="Path to the study plan file")
    parser.add_argument("--program", default="", help="Optional program hint")
    parser.add_argument("--transcript", default="", help="Optional transcript file path for comparison")
    parser.add_argument("--no-model-compare", action="store_true", help="Disable Gemini model comparison and use rule-based matching")
    args = parser.parse_args()

    if args.transcript:
        result = analyze_transcript_and_study_plan(
            transcript_file_path=args.transcript,
            study_plan_file_path=args.study_plan_file,
            program_hint=args.program or None,
            use_model_comparison=not args.no_model_compare,
        )
    else:
        result = analyze_study_plan(args.study_plan_file, program_hint=args.program or None)

    print(json.dumps(result, ensure_ascii=False, indent=2))
